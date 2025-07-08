import os
import re
import json
import uuid
import shutil
import asyncio
from pathlib import Path
from datetime import datetime
import flet as ft
import google.generativeai as genai
from dateutil import parser as dateparser

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not installed. .docx export disabled.")

# Session storage directory
SESSION_DIR = Path("sessions")
SESSION_DIR.mkdir(exist_ok=True)

# Custom dark theme
custom_theme = ft.Theme(
    color_scheme=ft.ColorScheme(
        primary=ft.Colors.BLUE,
        on_primary=ft.Colors.WHITE,
        surface=ft.Colors.GREY_900,
        on_surface=ft.Colors.WHITE,
        outline=ft.Colors.LIGHT_BLUE_200,
        outline_variant=ft.Colors.LIGHT_BLUE_100,
    ),
    use_material3=True
)

class ChatLogAnalyzer:
    """Handles API key, model listing, chat loading, person parsing, and analysis streaming."""
    FREE_TIER_LIMITS = {"default": 15, "Gemini 1.5 Flash": 15, "Gemini 1.5 Pro": 2}
    TIER_1_LIMITS   = {"default": 100, "Gemini 1.5 Flash": 1000, "Gemini 1.5 Pro": 10}

    def __init__(self):
        self.api_key            = None
        self.available_models   = []    # list of dicts {name, displayName}
        self.selected_model     = None  # internal model name
        self.chat_log_path      = None
        self.chat_log_content   = None
        self.persons            = []    # full unique list
        self.team1              = []    # Customer
        self.team2              = []    # IT Support
        self.llm_config         = {"base_instructions": "", "user_query": "", "language": "English"}
        self.is_analyzing       = False
        self.last_request_times = {}

    def load_api_key(self, key: str) -> bool:
        try:
            genai.configure(api_key=key)
            self.api_key = key
            return True
        except Exception as e:
            self.api_key = None
            print(f"API key configuration error: {e}")
            return False

    def fetch_available_models(self) -> tuple[bool,str,list]:
        """List all GenAI models supporting generateContent."""
        if not self.api_key:
            return False, "API key not set.", []
        try:
            models = [
                {"name": m.name, "displayName": m.display_name}
                for m in genai.list_models()
                if "generateContent" in m.supported_generation_methods
            ]
            # sort by displayName
            models.sort(key=lambda m: m["displayName"].lower())
            self.available_models = models
            return True, f"Found {len(models)} models.", models
        except Exception as e:
            return False, str(e), []

    def load_chat_log(self, path: str) -> bool:
        try:
            with open(path, "r", encoding="utf-8") as f:
                self.chat_log_content = f.read()
            self.chat_log_path = path
            return True
        except Exception as e:
            print(f"Error loading chat log: {e}")
            return False

    def get_date_range_from_log(self) -> tuple[str|None,str|None]:
        """Parse earliest and latest timestamps in WhatsApp export format."""
        if not self.chat_log_content:
            return None, None
        # e.g. "01/02/2025, 13:45" possibly with brackets
        pattern = re.compile(r"^\[?(\d{1,2}/\d{1,2}/\d{2,4}),\s*(\d{1,2}:\d{2}(?::\d{2})?)\]?", re.MULTILINE)
        dts = []
        for m in pattern.finditer(self.chat_log_content):
            try:
                dt = dateparser.parse(f"{m.group(1)} {m.group(2)}", dayfirst=True)
                dts.append(dt)
            except:
                continue
        if not dts:
            return None, None
        start, end = min(dts), max(dts)
        return start.strftime("%Y-%m-%d %H:%M"), end.strftime("%Y-%m-%d %H:%M")

    def parse_persons_with_llm(self, parser_name: str) -> tuple[bool,str]:
        """Use the LLM to extract unique participant names."""
        if not self.chat_log_content:
            return False, "Chat log not loaded."
        if not self.api_key:
            return False, "API key not set."
        if not parser_name:
            return False, "Parser model not selected."
        prompt = (
            "You are a WhatsApp chat log parser. Extract each unique participant name exactly as it appears "
            "in the log. List one per line, no additional text.\n\n"
            f"---BEGIN LOG---\n{self.chat_log_content}\n---END LOG---"
        )
        try:
            model = genai.GenerativeModel(parser_name)
            resp = model.generate_content(prompt)
            text = getattr(resp, "text", "")
            lines = [l.strip() for l in text.splitlines() if l.strip()]
            # Dedupe, but preserve order
            seen = []
            for ln in lines:
                if ln not in seen:
                    seen.append(ln)
            # update persons, preserving existing team assignments
            new_set = set(seen)
            self.persons = seen
            self.team1 = [p for p in self.team1 if p in new_set]
            self.team2 = [p for p in self.team2 if p in new_set]
            return True, f"Found {len(seen)} participants."
        except Exception as e:
            return False, f"LLM parsing error: {e}"

    def set_selected_model(self, display_name: str) -> bool:
        """Map displayName → internal model.name."""
        for m in self.available_models:
            if m["displayName"] == display_name:
                self.selected_model = m["name"]
                return True
        return False

    def is_rate_limited(self, model_disp: str, tier: str) -> tuple[bool,float]:
        """Simple RPM → seconds check."""
        limits = self.TIER_1_LIMITS if tier == "Tier 1" else self.FREE_TIER_LIMITS
        rpm = limits.get(model_disp, limits["default"])
        sec_per = 60.0 / rpm if rpm else 0
        last = self.last_request_times.get(model_disp)
        if last:
            elapsed = (datetime.now() - last).total_seconds()
            if elapsed < sec_per:
                return True, sec_per - elapsed
        return False, 0.0

    def record_request_time(self, model_disp: str):
        self.last_request_times[model_disp] = datetime.now()

    def translate_text_with_llm(self, text: str, target_lang: str, model_name: str) -> tuple[bool, str]:
        """Translate text to target_lang using the given model."""
        if not self.api_key:
            return False, "API key not set."
        if not text.strip():
            return False, "No text provided for translation."
        try:
            model = genai.GenerativeModel(model_name)
            prompt = (
                f"Translate the following text to {target_lang}. "
                "Return only the translated text without any additional commentary.\n\n"
                f"---\n{text}\n---"
            )
            resp = model.generate_content(prompt)
            translated = getattr(resp, "text", "").strip()
            if not translated:
                return False, "Empty translation result."
            return True, translated
        except Exception as e:
            return False, f"Translation error: {e}"

    def get_analysis_stream(self,
            start_dt: str,
            end_dt:   str,
            team1:    list[str],
            team2:    list[str],
            cfg:      dict
        ):
        """Generator yielding analysis chunks; statuses prefixed with 'Status:' go to Live Log."""
        if not self.selected_model or not self.chat_log_content:
            yield "Status: Error: model or chat log missing."
            return
        self.is_analyzing = True
        cust = ", ".join(team1) if team1 else "None"
        supp = ", ".join(team2) if team2 else "None"
        full_instr = (
            cfg["base_instructions"]
            + f"\n\nTeams → Customer: [{cust}], IT Support: [{supp}]\n"
        )
        prompt = (
            f"{full_instr}"
            f"Analyze chat between {start_dt} and {end_dt}.\n---\n{self.chat_log_content}\n---\n"
            f"User Query:\n{cfg['user_query']}"
        )
        try:
            yield "Status: Preparing prompt..."
            model = genai.GenerativeModel(self.selected_model)
            yield "Status: Sending to model..."
            stream = model.generate_content(prompt, stream=True)
            # wait for first real chunk to switch Tab
            first_chunk = True
            for chunk in stream:
                if not self.is_analyzing:
                    yield "\n--- ANALYSIS CANCELLED ---"
                    break
                text = getattr(chunk, "text", "")
                if text:
                    yield text
                elif first_chunk:
                    # no text at all
                    reason = getattr(chunk.candidates[0], "finish_reason", "UNKNOWN")
                    yield f"\n\n--- Model ended ({reason}) ---"
                first_chunk = False
        except Exception as e:
            yield f"\n\n--- ANALYSIS FAILED: {e}"
        finally:
            self.is_analyzing = False
class ChatAnalyzerApp:
    DEFAULT_PROMPT_ID = "default-detailed-report-v1"
    SESSION_FILE     = "session.json"

    def __init__(self, page: ft.Page):
        self.page            = page
        self.analyzer        = ChatLogAnalyzer()
        self.session         = self._load_latest_session()
        self.prompts         = []
        self.selected_persons = set()

        # Page setup
        page.title      = "WhatsApp Chat Analyzer"
        page.theme      = custom_theme
        page.theme_mode = ft.ThemeMode.DARK
        page.padding    = 10
        page.bgcolor    = ft.Colors.GREY_900
        page.on_window_event = self._on_window_event

        # Build UI
        self._setup_ui_components()
        self._build_layout()

        # Restore state
        page.run_task(self._load_state_from_session)

    def _setup_ui_components(self):
        common = {"color": ft.Colors.WHITE, "border_color": ft.Colors.LIGHT_BLUE_200}
        # --- Step 1: API Key & Tier ---
        self.api_key_field = ft.TextField(
            label="API Key",
            password=True,
            can_reveal_password=True,
            value=self.session.get("api_key", ""),
            on_submit=self._on_save_api_key_and_fetch,
            expand=True,
            **common
        )
        self.tier_dropdown = ft.Dropdown(
            label="API Tier",
            options=[
                ft.dropdown.Option("Free Tier"),
                ft.dropdown.Option("Tier 1")
            ],
            value=self.session.get("api_tier", "Free Tier"),
            **common
        )
        self.save_api_btn = ft.ElevatedButton(
            "Save & Fetch Models",
            icon=ft.Icons.SAVE,
            on_click=self._on_save_api_key_and_fetch
        )

        # --- Step 2: Load & Parse Chat Log ---
        self.chat_log_label   = ft.Text(self.session.get("chat_log_display", "No chat log loaded"), expand=True, no_wrap=True)
        self.load_log_btn     = ft.ElevatedButton("Load Chat Log", icon=ft.Icons.ATTACH_FILE, on_click=self._on_load_chat_log)
        self.parser_model_dd  = ft.Dropdown(label="Parser Model", options=[], disabled=True, expand=True, **common)
        self.parse_btn        = ft.ElevatedButton("Parse Persons", icon=ft.Icons.PLAY_ARROW, disabled=True, on_click=lambda e: self.page.run_task(self._on_parse_persons))
        self.stop_parse_btn   = ft.ElevatedButton("Stop", icon=ft.Icons.CANCEL, style=ft.ButtonStyle(bgcolor=ft.Colors.RED_400), visible=False, on_click=self._on_stop_parse)

        # --- Step 3: Assign Teams ---
        self.assign_cust_btn    = ft.ElevatedButton("Assign → Customer", disabled=True, on_click=self._on_assign_customer)
        self.assign_support_btn = ft.ElevatedButton("Assign → IT Support", disabled=True, on_click=self._on_assign_support)
        self.unassign_btn       = ft.ElevatedButton("Unassign ←", disabled=True, on_click=self._on_unassign)
        self.unassigned_grid    = ft.Row([], spacing=8, wrap=True)
        self.grid_cust          = ft.Row([], spacing=8, wrap=True)
        self.grid_support       = ft.Row([], spacing=8, wrap=True)

        # --- Step 4: Date/Time Range ---
        self.start_date_field = ft.TextField(label="Start Date", read_only=True, expand=True, **common)
        self.start_time_field = ft.TextField(label="Start Time", read_only=True, width=120, **common)
        self.end_date_field   = ft.TextField(label="End Date", read_only=True, expand=True, **common)
        self.end_time_field   = ft.TextField(label="End Time", read_only=True, width=120, **common)

        # Date/Time pickers in overlay
        self.dp_start = ft.DatePicker(on_change=self._on_date_change)
        self.tp_start = ft.TimePicker(on_change=self._on_date_change)
        self.dp_end   = ft.DatePicker(on_change=self._on_date_change)
        self.tp_end   = ft.TimePicker(on_change=self._on_date_change)
        self.page.overlay.extend([self.dp_start, self.tp_start, self.dp_end, self.tp_end])

        # --- Live Log ---
        self.live_log = ft.ListView(spacing=3, auto_scroll=True, expand=True)

        # --- Tab 2: Prompts & Analysis ---
        self.prompt_list         = ft.ListView(spacing=5, auto_scroll=True, expand=True)
        self.add_prompt_btn      = ft.ElevatedButton("Add", icon=ft.Icons.ADD, on_click=lambda e: self._edit_prompt(None))
        self.import_prompts_btn  = ft.IconButton(icon=ft.Icons.FILE_UPLOAD, on_click=self._on_import_prompts, tooltip="Import prompts")
        self.export_prompts_btn  = ft.IconButton(icon=ft.Icons.FILE_DOWNLOAD, on_click=self._on_export_prompts, tooltip="Export prompts")
        self.prompt_name_field   = ft.TextField(label="Prompt Name", **common)
        self.base_instr_field    = ft.TextField(label="Base Instructions", multiline=True, min_lines=5, max_lines=7, **common)
        self.user_query_field    = ft.TextField(label="User Query", multiline=True, min_lines=5, max_lines=7, **common)
        self.analysis_model_dd   = ft.Dropdown(label="Analysis Model", options=[], on_change=self._on_analysis_model_select, expand=True, **common)
        self.language_field      = ft.TextField(label="Language", value="English", expand=True, tooltip="If not English, query will be auto-translated", **common)
        self.save_prompt_btn     = ft.ElevatedButton("Save Prompt", icon=ft.Icons.SAVE, on_click=self._on_save_prompt)
        self.delete_prompt_btn   = ft.ElevatedButton("Delete Prompt", icon=ft.Icons.DELETE, style=ft.ButtonStyle(color=ft.Colors.RED), on_click=self._on_delete_prompt, disabled=True)
        self.run_analysis_btn    = ft.ElevatedButton("Run Analysis", icon=ft.Icons.PLAY_ARROW, on_click=self._on_run_analysis)
        self.stop_analysis_btn   = ft.ElevatedButton("Stop", icon=ft.Icons.CANCEL, style=ft.ButtonStyle(bgcolor=ft.Colors.RED_200), visible=False, on_click=self._on_stop_analysis)
        self.timer_text          = ft.Text("0.0s", weight=ft.FontWeight.BOLD)
        self.thinking_indicator  = ft.Row([ft.ProgressRing(), ft.Text("Thinking...")], spacing=10, visible=False)
        self.generating_indicator= ft.Row([ft.ProgressRing(), ft.Text("Generating...")], spacing=10, visible=False)

        # --- Tab 3: Results ---
        self.results_md          = ft.Markdown("Results will appear here.", expand=True, selectable=True)
        self.copy_results_btn    = ft.ElevatedButton("Copy", icon=ft.Icons.COPY, on_click=self._on_copy_results)
        self.save_results_btn    = ft.ElevatedButton("Save As…", icon=ft.Icons.SAVE_AS, on_click=self._on_save_results)
    def _build_layout(self):
        # Step 1+2 container
        step12 = ft.Column([
            ft.Text("1. Configure API", style=ft.TextThemeStyle.HEADLINE_SMALL),
            ft.Row([self.api_key_field, self.tier_dropdown, self.save_api_btn], spacing=8),
            ft.Divider(),
            ft.Text("2. Load & Parse Chat Log", style=ft.TextThemeStyle.HEADLINE_SMALL),
            ft.Row([self.load_log_btn, self.chat_log_label], spacing=8, vertical_alignment=ft.CrossAxisAlignment.CENTER),
            ft.Row([self.parser_model_dd, self.parse_btn, self.stop_parse_btn], spacing=8),
        ], spacing=10, expand=False)

        # Step 3+4 container
        step34 = ft.Column([
            ft.Text("3. Assign Teams", style=ft.TextThemeStyle.HEADLINE_SMALL),
            ft.Text("Select persons → use buttons", color=ft.Colors.LIGHT_BLUE_100),
            self.unassigned_grid,
            ft.Row([self.assign_cust_btn, self.assign_support_btn, self.unassign_btn], spacing=8),
            ft.Row([
                ft.Column([ft.Text("Customer Team"), self.grid_cust], expand=1),
                ft.VerticalDivider(),
                ft.Column([ft.Text("IT Support Team"), self.grid_support], expand=1),
            ], expand=True),
            ft.Divider(),
            ft.Text("4. Confirm Date Range", style=ft.TextThemeStyle.HEADLINE_SMALL),
            ft.Row([
                self.start_date_field,
                ft.IconButton(ft.Icons.CALENDAR_MONTH, tooltip="Pick Start Date", on_click=lambda e: self._open_picker(self.dp_start)),
                self.start_time_field,
                ft.IconButton(ft.Icons.ACCESS_TIME, tooltip="Pick Start Time", on_click=lambda e: self._open_picker(self.tp_start)),
            ], spacing=8),
            ft.Row([
                self.end_date_field,
                ft.IconButton(ft.Icons.CALENDAR_MONTH, tooltip="Pick End Date", on_click=lambda e: self._open_picker(self.dp_end)),
                self.end_time_field,
                ft.IconButton(ft.Icons.ACCESS_TIME, tooltip="Pick End Time", on_click=lambda e: self._open_picker(self.tp_end)),
            ], spacing=8),
        ], spacing=10, expand=True)

        # Live Log container
        live_log_cont = ft.Container(self.live_log, border=ft.border.all(1, ft.Colors.LIGHT_BLUE_200), padding=8, height=150)

        # Tab 1: Setup

        setup_tab = ft.Tab(
                text="1. Setup", icon=ft.Icons.SETTINGS,
                content=ft.Container(
                 ft.Column(
                    [ step12, ft.Divider(), step34 ],
                    scroll=ft.ScrollMode.ADAPTIVE,  # ← move scroll here
                    expand=True
                ),
                padding=15
            )
        )
 

        # Tab 2: Prompts & Analysis
        analysis_left = ft.Column([
            ft.Text("Prompt Library", style=ft.TextThemeStyle.HEADLINE_SMALL),
            ft.Row([self.add_prompt_btn, self.import_prompts_btn, self.export_prompts_btn], spacing=8),
            ft.Divider(),
            self.prompt_list
        ], expand=1, spacing=10)

        analysis_right = ft.Column([
            ft.Text("Prompt Editor", style=ft.TextThemeStyle.HEADLINE_SMALL),
            self.prompt_name_field,
            ft.Row([self.analysis_model_dd, self.language_field], spacing=8),
            self.base_instr_field,
            self.user_query_field,
            ft.Row([self.save_prompt_btn, self.delete_prompt_btn], spacing=8),
            ft.Divider(),
            ft.Row([ft.Stack([self.run_analysis_btn, self.stop_analysis_btn]), self.thinking_indicator, self.generating_indicator, self.timer_text], spacing=20)
        ], expand=2, scroll=ft.ScrollMode.ADAPTIVE, spacing=10)

        analysis_tab = ft.Tab(
            text="2. Prompts & Analysis", icon=ft.Icons.PLAYLIST_ADD_CHECK,
            content=ft.Container(ft.Row([analysis_left, ft.VerticalDivider(), analysis_right], expand=True), padding=15)
        )

        # Tab 3: Results
        results_tab = ft.Tab(
            text="3. Results", icon=ft.Icons.ARTICLE,
            content=ft.Container(
                ft.Column([
                    ft.Row([self.copy_results_btn, self.save_results_btn], alignment=ft.MainAxisAlignment.END),
                    ft.Text("Save as Word (.docx) for formatted export.", size=12, italic=True),
                    self.results_md
                ], spacing=10),
                padding=15
            )
        )

        # Assemble Tabs + Live Log
        self.tabs = ft.Tabs(selected_index=0, expand=True, tabs=[setup_tab, analysis_tab, results_tab])
        self.page.add(ft.Column([self.tabs, live_log_cont], expand=True, spacing=10))

    # --- Utility to open a Date/Time picker ---
    def _open_picker(self, picker: ft.Control):
        picker.open = True
        self.page.update()

    # --- Logging helper ---
    def log(self, msg: str, icon="ℹ️"):
        ts = datetime.now().strftime("[%H:%M:%S]")
        self.live_log.controls.append(ft.Text(f"{ts} {icon} {msg}", size=11, color=ft.Colors.WHITE))
        if len(self.live_log.controls) > 200:
            self.live_log.controls.pop(0)
        self.page.update()

    def log_message(self, msg: str, icon="ℹ️"):
        self.log(msg, icon)
    # -----------------
    # Session I/O
    # -----------------
    def _load_latest_session(self) -> dict:
        files = sorted(SESSION_DIR.glob("session_*.json"), reverse=True)
        if not files:
            return {}
        try:
            return json.loads(files[0].read_text(encoding="utf-8"))
        except:
            return {}

    def _save_session(self):
        data = {
            "api_key":        self.api_key_field.value,
            "api_tier":       self.tier_dropdown.value,
            "parser_model":   self.parser_model_dd.value,
            "chat_log_display": self.chat_log_label.value,
            "chat_log_path":  self.analyzer.chat_log_path,
            "start_iso":      self.session.get("start_iso"),
            "end_iso":        self.session.get("end_iso"),
            "model_display":  self.analysis_model_dd.value,
            "team1":          self.analyzer.team1,
            "team2":          self.analyzer.team2
        }
        fn = SESSION_DIR / f"session_{datetime.now():%Y%m%d_%H%M%S}.json"
        fn.write_text(json.dumps(data, indent=2), encoding="utf-8")
        self.log("Session saved", "💾")

    def _on_window_event(self, e):
        if e.data == "close":
            self._save_session()
            self.page.window_destroy()

    # -----------------
    # Load State
    # -----------------
    async def _load_state_from_session(self):
        # prompts
        self._load_prompts()
        self._refresh_prompt_list()

        key = self.session.get("api_key")
        if key and self.analyzer.load_api_key(key):
            self.api_key_field.value = key
            await self._fetch_models()
        path = self.session.get("chat_log_path")
        if path and os.path.exists(path):
            await self._load_chat_log(path)
            # auto-parse-and-refresh
            await self._on_parse_persons()
            self.analyzer.team1 = self.session.get("team1", [])
            self.analyzer.team2 = self.session.get("team2", [])
            self._refresh_assignment_buttons()
            self._refresh_grids()

        if self.session.get("start_iso"):
            # replicate on_date_change logic
            dt = dateparser.parse(self.session["start_iso"])
            self.dp_start.value, self.tp_start.value = dt.date(), dt.time()
            self._on_date_change(None)
        if self.session.get("end_iso"):
            dt = dateparser.parse(self.session["end_iso"])
            self.dp_end.value, self.tp_end.value = dt.date(), dt.time()
            self._on_date_change(None)

        self.page.update()

    # -----------------
    # API Key & Models
    # -----------------
    async def _on_save_api_key_and_fetch(self, e=None):
        key = self.api_key_field.value.strip()
        if not key:
            self.log_message("Enter an API key first", "🛑"); return
        self.log_message("Configuring API key...", "🔑")
        if self.analyzer.load_api_key(key):
            self.log_message("API key set.", "✅")
            await self._fetch_models()
            self._save_session()
        else:
            self.log_message("Invalid API key", "🛑")

    async def _fetch_models(self):
        self.parser_model_dd.disabled = True
        self.analysis_model_dd.disabled = True
        ok, msg, models = await asyncio.to_thread(self.analyzer.fetch_available_models)
        if ok:
            opts = [ft.dropdown.Option(m["displayName"]) for m in models]
            self.parser_model_dd.options  = opts.copy()
            self.analysis_model_dd.options = opts.copy()
            # restore last
            pm = self.session.get("parser_model")
            am = self.session.get("model_display")
            self.parser_model_dd.value   = pm if pm in [o.key for o in opts] else opts[0].key if opts else None
            self.analysis_model_dd.value = am if am in [o.key for o in opts] else opts[0].key if opts else None
            self.log_message(f"Loaded {len(models)} models", "📡")
        else:
            self.log_message(f"Model fetch failed: {msg}", "🛑")
        self.parser_model_dd.disabled  = False
        self.analysis_model_dd.disabled= False
        self.page.update()

    def _on_analysis_model_select(self, e):
        val = self.analysis_model_dd.value
        if val and self.analyzer.set_selected_model(val):
            self.log_message(f"Analysis model set to {val}", "🤖")
        else:
            self.log_message(f"Failed to set analysis model: {val}", "🛑")

    # -----------------
    # Load / Parse Chat Log
    # -----------------
    async def _on_load_chat_log(self, e):
        dlg = ft.FilePicker(on_result=self._on_load_chat_log_result)
        self.page.overlay.append(dlg)
        self.page.update()
        dlg.pick_files(allow_multiple=False, allowed_extensions=["txt"])

    async def _on_load_chat_log_result(self, e: ft.FilePickerResultEvent):
        if not e.files:
            self.log_message("Load cancelled", "🚫"); return
        await self._load_chat_log(e.files[0].path)

    async def _load_chat_log(self, path: str):
        self.log("Loading chat log...", "📄")
        ok = await asyncio.to_thread(self.analyzer.load_chat_log, path)
        if not ok:
            self.log_message("Failed to load chat log", "🛑"); return
        fname = os.path.basename(path)
        self.chat_log_label.value = f"Loaded: {fname}"
        self.session["chat_log_display"] = self.chat_log_label.value
        start, end = await asyncio.to_thread(self.analyzer.get_date_range_from_log)
        if start and end:
            dt_s = dateparser.parse(start)
            dt_e = dateparser.parse(end)
            self.dp_start.value, self.tp_start.value = dt_s.date(), dt_s.time()
            self.dp_end.value,   self.tp_end.value   = dt_e.date(), dt_e.time()
            self._on_date_change(None)
        self.parse_btn.disabled = False
        self.log_message("Chat log loaded", "✅")
        self.page.update()

    async def _on_parse_persons(self, e=None):
        if not (self.analyzer.chat_log_content and self.parser_model_dd.value):
            self.log_message("Missing chat log or parser model", "🛑"); return
        disp = self.parser_model_dd.value
        self.log(f"Parsing persons with {disp}...", "▶️")
        self.parse_btn.disabled = True
        self.stop_parse_btn.visible = True
        self.page.update()
        name = next((m["name"] for m in self.analyzer.available_models if m["displayName"] == disp), None)
        ok, msg = await asyncio.to_thread(self.analyzer.parse_persons_with_llm, name)
        self.log_message(msg, "✅" if ok else "🛑")
        self.selected_persons.clear()
        if ok:
            self._refresh_grids()
        self.parse_btn.disabled = False
        self.stop_parse_btn.visible = False
        self._refresh_assignment_buttons()
        self._save_session()
        self.page.update()

    def _on_stop_parse(self, e):
        # purely cosmetic—no cancellation in backend
        self.log_message("Parse stop requested", "🛑")
        self.parse_btn.disabled = False
        self.stop_parse_btn.visible = False
        self.page.update()

    # -----------------
    # Team Assignment
    # -----------------
    def _toggle_person(self, name: str):
        if name in self.selected_persons:
            self.selected_persons.remove(name)
        else:
            self.selected_persons.add(name)
        self._refresh_assignment_buttons()
        self._refresh_grids()

    def _refresh_assignment_buttons(self):
        any_sel = bool(self.selected_persons)
        for btn in (self.assign_cust_btn, self.assign_support_btn, self.unassign_btn):
            btn.disabled = not any_sel
        self.page.update()

    def _on_assign_customer(self, e):
        for p in list(self.selected_persons):
            if p in self.analyzer.team2:
                self.analyzer.team2.remove(p)
            if p not in self.analyzer.team1:
                self.analyzer.team1.append(p)
        self.log_message(f"Assigned {len(self.selected_persons)} to Customer", "👤")
        self.selected_persons.clear()
        self._refresh_assignment_buttons()
        self._refresh_grids()
        self._save_session()

    def _on_assign_support(self, e):
        for p in list(self.selected_persons):
            if p in self.analyzer.team1:
                self.analyzer.team1.remove(p)
            if p not in self.analyzer.team2:
                self.analyzer.team2.append(p)
        self.log_message(f"Assigned {len(self.selected_persons)} to IT Support", "🛠️")
        self.selected_persons.clear()
        self._refresh_assignment_buttons()
        self._refresh_grids()
        self._save_session()

    def _on_unassign(self, e):
        for p in list(self.selected_persons):
            if p in self.analyzer.team1:
                self.analyzer.team1.remove(p)
            if p in self.analyzer.team2:
                self.analyzer.team2.remove(p)
        self.log_message(f"Unassigned {len(self.selected_persons)} person(s)", "♻️")
        self.selected_persons.clear()
        self._refresh_assignment_buttons()
        self._refresh_grids()
        self._save_session()

    def _refresh_grids(self):
        def btn(p):
            sel = p in self.selected_persons
            return ft.ElevatedButton(
                p,
                on_click=lambda e, x=p: self._toggle_person(x),
                bgcolor=ft.Colors.LIGHT_BLUE_200 if sel else None,
                color=ft.Colors.BLACK if sel else ft.Colors.WHITE,
                style=ft.ButtonStyle(shape=ft.RoundedRectangleBorder(radius=8))
            )
        all_persons = self.analyzer.persons
        unassigned = [p for p in all_persons if p not in self.analyzer.team1 and p not in self.analyzer.team2]
        self.unassigned_grid.controls = [btn(p) for p in sorted(unassigned)]
        self.grid_cust.controls     = [btn(p) for p in sorted(self.analyzer.team1)]
        self.grid_support.controls  = [btn(p) for p in sorted(self.analyzer.team2)]
        self.page.update()

    # -----------------
    # Date / Time Handling
    # -----------------
    def _on_date_change(self, e):
        if self.dp_start.value:
            self.start_date_field.value = self.dp_start.value.strftime("%Y-%m-%d")
        if self.tp_start.value:
            self.start_time_field.value = self.tp_start.value.strftime("%H:%M")
        if self.dp_end.value:
            self.end_date_field.value = self.dp_end.value.strftime("%Y-%m-%d")
        if self.tp_end.value:
            self.end_time_field.value = self.tp_end.value.strftime("%H:%M")
        try:
            if self.dp_start.value and self.tp_start.value:
                self.session["start_iso"] = datetime.combine(self.dp_start.value, self.tp_start.value).isoformat()
            if self.dp_end.value and self.tp_end.value:
                self.session["end_iso"] = datetime.combine(self.dp_end.value, self.tp_end.value).isoformat()
            self._save_session()
        except Exception:
            pass
        self.page.update()

    # -----------------
    # Prompts Library CRUD
    # -----------------
    def _load_prompts(self):
        try:
            with open("prompts.json","r",encoding="utf-8") as f:
                self.prompts = json.load(f)
        except:
            self.prompts = []
        if not any(p.get("id")==self.DEFAULT_PROMPT_ID for p in self.prompts):
            # insert default prompt
            default = {
                "id": self.DEFAULT_PROMPT_ID,
                "name": "Default IT Support Report",
                "model": "Gemini 1.5 Pro",
                "language": "English",
                "base_instructions": "...",
                "user_query":   "Analyze the chat log..."
            }
            self.prompts.insert(0, default)

    def _refresh_prompt_list(self):
        self.prompt_list.controls.clear()
        for p in sorted(self.prompts, key=lambda x: x["name"].lower()):
            icon = ft.Icons.STAR if p["id"]==self.DEFAULT_PROMPT_ID else ft.Icons.DESCRIPTION
            self.prompt_list.controls.append(
                ft.ListTile(
                    leading=ft.Icon(icon),
                    title=ft.Text(p["name"]),
                    on_click=lambda e, pp=p: self._edit_prompt(pp)
                )
            )
        self.page.update()

    def _edit_prompt(self, data=None):
        if data:
            self.selected_prompt_id         = data["id"]
            self.prompt_name_field.value   = data["name"]
            self.analysis_model_dd.value   = data["model"]
            self.language_field.value      = data["language"]
            self.base_instr_field.value    = data["base_instructions"]
            self.user_query_field.value    = data["user_query"]
            self.delete_prompt_btn.disabled= (data["id"]==self.DEFAULT_PROMPT_ID)
            # set analyzer model
            self.analyzer.set_selected_model(data["model"])
        else:
            self.selected_prompt_id = None
            self.prompt_name_field.value = ""
            self.prompt_name_field.hint_text = "Enter prompt name"
            self.analysis_model_dd.value = self.analysis_model_dd.options[0].key if self.analysis_model_dd.options else None
            self.language_field.value      = "English"
            self.base_instr_field.value    = ""
            self.user_query_field.value    = ""
            self.delete_prompt_btn.disabled= True
        self.page.update()

    def _on_save_prompt(self, e):
        name = self.prompt_name_field.value.strip()
        if not name:
            self.log_message("Prompt name cannot be empty", "🛑")
            return
        new = {
            "id":   self.selected_prompt_id or str(uuid.uuid4()),
            "name": name,
            "model": self.analysis_model_dd.value,
            "language": self.language_field.value,
            "base_instructions": self.base_instr_field.value,
            "user_query": self.user_query_field.value
        }
        idx = next((i for i,p in enumerate(self.prompts) if p["id"]==self.selected_prompt_id), -1)
        if idx>=0:
            self.prompts[idx] = new
        else:
            self.prompts.append(new)
            self.selected_prompt_id = new["id"]
        with open("prompts.json","w",encoding="utf-8") as f:
            json.dump(self.prompts,f,indent=4)
        self.log_message(f"Saved prompt '{name}'", "✅")
        self._refresh_prompt_list()

    def _on_delete_prompt(self, e):
        if self.selected_prompt_id==self.DEFAULT_PROMPT_ID:
            return
        self.prompts = [p for p in self.prompts if p["id"]!=self.selected_prompt_id]
        with open("prompts.json","w",encoding="utf-8") as f:
            json.dump(self.prompts,f,indent=4)
        self.log_message("Prompt deleted", "🗑️")
        self._refresh_prompt_list()
        self._edit_prompt(self.prompts[0] if self.prompts else None)

    async def _on_import_prompts(self, e):
        dlg = ft.FilePicker(on_result=self._on_import_prompts_result)
        self.page.overlay.append(dlg); self.page.update()
        dlg.pick_files(allow_multiple=False, allowed_extensions=["json"])

    async def _on_import_prompts_result(self, e: ft.FilePickerResultEvent):
        if not e.files:
            self.log_message("Import cancelled", "🚫")
            return
        try:
            shutil.copy2(e.files[0].path, "prompts.json")
            self._load_prompts(); self._refresh_prompt_list()
            self.log_message("Prompts imported", "✅")
        except Exception as ex:
            self.log_message(f"Import error: {ex}", "🛑")

    async def _on_export_prompts(self, e):
        dlg = ft.FilePicker(on_result=self._on_export_prompts_result)
        self.page.overlay.append(dlg); self.page.update()
        dlg.save_file(file_name="prompts_export.json", allowed_extensions=["json"])

    async def _on_export_prompts_result(self, e: ft.FilePickerResultEvent):
        if not e.path:
            self.log_message("Export cancelled", "🚫")
            return
        try:
            shutil.copy2("prompts.json", e.path)
            self.log_message("Prompts exported", "✅")
        except Exception as ex:
            self.log_message(f"Export error: {ex}", "🛑")

    # -----------------
    # Analysis
    # -----------------
    def _on_run_analysis(self, e):
        if self.analyzer.is_analyzing:
            self.log_message("Already analyzing...", "⏳"); return
        if not self.analyzer.chat_log_content:
            self.log_message("Load chat log first.", "🛑"); return
        if not self.analyzer.selected_model:
            self.log_message("Select an analysis model", "🛑"); return
        rl, wait = self.analyzer.is_rate_limited(self.analysis_model_dd.value, self.tier_dropdown.value)
        if rl:
            self.log_message(f"Rate limit: wait {wait:.1f}s", "🛑")
            return
        # clear results
        self.results_md.value = ""
        # show running indicators
        self.run_analysis_btn.visible = False
        self.stop_analysis_btn.visible= True
        self.thinking_indicator.visible   = True
        self.generating_indicator.visible = False
        self.timer_text.value = "0.0s"
        self.tabs.selected_index = 2
        self.page.update()
        self.page.run_task(self._run_analysis_task)

    async def _run_analysis_task(self):
        # optional translation
        uq = self.user_query_field.value
        if self.language_field.value.lower() != "english":
            self.log_message("Translating query to English...", "🌐")
            ok, tr = await asyncio.to_thread(
                self.analyzer.translate_text_with_llm,
                uq, "English", self.analyzer.selected_model
            )
            if not ok:
                self.log_message(tr, "🛑")
                self._reset_analysis_controls()
                return
            uq = tr
            self.log_message("Translation complete", "✅")
        # configure
        self.analyzer.llm_config = {
            "base_instructions": self.base_instr_field.value,
            "user_query": uq,
            "language": "English"
        }
        self.analyzer.record_request_time(self.analysis_model_dd.value)
        start = datetime.now()
        timer = asyncio.create_task(self._update_timer(start))

        # stream
        first_chunk = True
        for chunk in self.analyzer.get_analysis_stream(
                self.start_date_field.value or "N/A",
                self.end_date_field.value   or "N/A",
                self.analyzer.team1,
                self.analyzer.team2,
                self.analyzer.llm_config
            ):
            if chunk.startswith("Status:"):
                self.log_message(chunk[len("Status:"):].strip(), "📡")
                if first_chunk:
                    # now switch to generating
                    self.thinking_indicator.visible   = False
                    self.generating_indicator.visible = True
                    first_chunk = False
                self.page.update()
                continue
            # actual content
            self.results_md.value += chunk
            self.page.update()
            await asyncio.sleep(0.01)

        # done
        self.analyzer.is_analyzing = False
        await timer
        self._reset_analysis_controls()
        self.log_message("Analysis complete", "✅")

    async def _update_timer(self, start):
        while self.analyzer.is_analyzing:
            elapsed = (datetime.now() - start).total_seconds()
            self.timer_text.value = f"{elapsed:.1f}s"
            self.page.update()
            await asyncio.sleep(0.1)

    def _on_stop_analysis(self, e):
        self.analyzer.is_analyzing = False
        self.log_message("Analysis cancelled", "🛑")

    def _reset_analysis_controls(self):
        self.run_analysis_btn.visible     = True
        self.stop_analysis_btn.visible    = False
        self.thinking_indicator.visible   = False
        self.generating_indicator.visible = False
        self.page.update()

    def _on_copy_results(self, e):
        if self.results_md.value:
            self.page.set_clipboard(self.results_md.value)
            self.log_message("Results copied", "✅")

    def _on_save_results(self, e):
        dlg = ft.FilePicker(on_result=self._on_save_analysis_result)
        self.page.overlay.append(dlg); self.page.update()
        exts = ["docx","txt"] if DOCX_AVAILABLE else ["txt"]
        dlg.save_file(dialog_title="Save Analysis", file_name="analysis.docx", allowed_extensions=exts)

    async def _on_save_analysis_result(self, e: ft.FilePickerResultEvent):
        if not e.path:
            self.log_message("Save cancelled", "🚫"); return
        path = Path(e.path)
        try:
            if path.suffix == ".docx" and DOCX_AVAILABLE:
                doc = Document()
                doc.add_heading("Chat Analysis Report",0)
                for ln in self.results_md.value.splitlines():
                    clean = re.sub(r'\*\*|\*','',ln).strip()
                    if not clean: continue
                    if re.match(r'^\d+\.\d+\.\d+', clean):
                        doc.add_paragraph(clean, style='Heading 3')
                    elif re.match(r'^\d+\.\d+', clean):
                        doc.add_paragraph(clean, style='Heading 2')
                    elif re.match(r'^\d+\.', clean):
                        doc.add_paragraph(clean, style='Heading 1')
                    elif clean.startswith('-'):
                        doc.add_paragraph(clean, style='List Bullet')
                    else:
                        doc.add_paragraph(clean)
                doc.save(path)
            else:
                path.write_text(self.results_md.value, encoding="utf-8")
            self.log_message(f"Saved results to {path.name}", "✅")
        except Exception as ex:
            self.log_message(f"Save failed: {ex}", "🛑")
def main(page: ft.Page):
    page.window_min_width  = 1200
    page.window_min_height = 800
    ChatAnalyzerApp(page)

if __name__ == "__main__":
    ft.app(target=main)
